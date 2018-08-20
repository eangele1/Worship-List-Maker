#imports required modules to make document write-up functional
import calendar
import datetime
import csv
import random
import sys
import time

#tries to import docx; if docx is not installed, prompts user to run setup.py first and ends program
try:
	import docx
except ImportError:
	print "Python-docx is not installed. Run setup.py first."
	sys.exit(1)
#end of try-except

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#gets current year and month number
yearNum = datetime.datetime.now()
monthNum = datetime.datetime.now()

#gets third tuesday of month for leading prayer
tuesdayPrayer = calendar.Calendar(1).monthdatescalendar(yearNum.year, monthNum.month)[3][0]
tuesdayDate = "{:%d}".format(tuesdayPrayer)

#gets each wednesday of current month
wedFlag = False

#tries to see if Wednesday is the first day of the month; if not, makes first Wednesday date 0
try:
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[0][0]
	wednesdayDate1 = "{:%d}".format(wednesdayWorship)

	##DEBUG ONLY
	#print("TRY")
	#print(wednesdayDate1)
	##DEBUG ONLY

except IndexError:
	wednesdayDate1 = 0

	##DEBUG ONLY
	#print("EXCEPT")
	#print(wednesdayDate1)
	##DEBUG ONLY

#end

if int(wednesdayDate1) == 0 or int(wednesdayDate1) > 21:

	wedFlag = True

	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[1][0]
	wednesdayDate1 = "{:%d}".format(wednesdayWorship)
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[2][0]
	wednesdayDate2 = "{:%d}".format(wednesdayWorship)
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[3][0]
	wednesdayDate3 = "{:%d}".format(wednesdayWorship)
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[4][0]
	wednesdayDate4 = "{:%d}".format(wednesdayWorship)
	try:
		wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[5][0]
		wednesdayDate5 = "{:%d}".format(wednesdayWorship)

		##DEBUG ONLY
		#print("TRY")
		#print(wednesdayDate5)
		##DEBUG ONLY

	except IndexError:
		wednesdayDate5 = 0

		##DEBUG ONLY
		#print("EXCEPT")
		#print(wednesdayDate5)
		##DEBUG ONLY
		
	#end
else:
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[1][0]
	wednesdayDate2 = "{:%d}".format(wednesdayWorship)
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[2][0]
	wednesdayDate3 = "{:%d}".format(wednesdayWorship)
	wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[3][0]
	wednesdayDate4 = "{:%d}".format(wednesdayWorship)
	try:
		wednesdayWorship = calendar.Calendar(2).monthdatescalendar(yearNum.year, monthNum.month)[4][0]
		wednesdayDate5 = "{:%d}".format(wednesdayWorship)

		##DEBUG ONLY
		#print("TRY")
		#print(wednesdayDate5)
		##DEBUG ONLY

	except IndexError:
		wednesdayDate5 = 0

		##DEBUG ONLY
		#print("EXCEPT")
		#print(wednesdayDate5)
		##DEBUG ONLY

#end

#gets each sunday of current month

sunFlag = False

try:
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[0][0]
	sundayDate1 = "{:%d}".format(sundayWorship)

	##DEBUG ONLY
	#print("TRY")
	#print(sundayDate1)
	##DEBUG ONLY
	
except IndexError:
	sundayDate1 = 0

	##DEBUG ONLY
	#print("EXCEPT")
	#print(sundayDate1)
	##DEBUG ONLY
	
#end

if int(sundayDate1) == 0 or int(sundayDate1) > 23:

	sunFlag = True

	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[1][0]
	sundayDate1 = "{:%d}".format(sundayWorship)
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[2][0]
	sundayDate2 = "{:%d}".format(sundayWorship)
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[3][0]
	sundayDate3 = "{:%d}".format(sundayWorship)
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[4][0]
	sundayDate4 = "{:%d}".format(sundayWorship)
	try:
		sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[5][0]
		sundayDate5 = "{:%d}".format(sundayWorship)

		##DEBUG ONLY
		#print("SUN IN TRY")
		#print(sundayDate5)
		##DEBUG ONLY

	except IndexError:
		sundayDate5 = 0

		##DEBUG ONLY
		#print("SUN IN EXCEPT")
		#print(sundayDate5)
		##DEBUG ONLY

	#end
else:
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[1][0]
	sundayDate2 = "{:%d}".format(sundayWorship)
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[2][0]
	sundayDate3 = "{:%d}".format(sundayWorship)
	sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[3][0]
	sundayDate4 = "{:%d}".format(sundayWorship)
	try:
		sundayWorship = calendar.Calendar(6).monthdatescalendar(yearNum.year, monthNum.month)[4][0]
		sundayDate5 = "{:%d}".format(sundayWorship)

		##DEBUG ONLY
		#print("SUN OUT TRY")
		#print(sundayDate5)
		##DEBUG ONLY

	except IndexError:
		sundayDate5 = 0

		##DEBUG ONLY
		#print("SUN OUT EXCEPT")
		#print(sundayDate5)
		##DEBUG ONLY

#end

#months array
months = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

#information arrays

shirtColors = {'colors': None}

tempVestuario = []

HimnoWed = []
DirijeWed = []
BajoWed = []
GuitarraWed = []
PianoWed = []
BateriaWed = []
ApoyoWed = []

HimnoSun12 = []
DirijeSun12 = []
BajoSun12 = []
GuitarraSun12 = []
PianoSun12 = []
BateriaSun12 = []
ApoyoSun12 = []

HimnoSun6 = []
DirijeSun6 = []
BajoSun6 = []
GuitarraSun6 = []
PianoSun6 = []
BateriaSun6 = []
ApoyoSun6 = []

#now processing the information from CSV file

#opens input.csv and creates object
try:
	f = open('input.csv')
except IOError:
	print "The input.csv is not inside the same directory as the program. Please place the file inside this directory."
	sys.exit(1)
#end
people = csv.reader(f, quotechar = '"', delimiter = ',', quoting = csv.QUOTE_ALL, skipinitialspace = True)

#creates newNumbers.csv and creates object
csvfile = open('newNumbers.csv', 'wb')
filewriter = csv.writer(csvfile, delimiter=',', quotechar='|', quoting=csv.QUOTE_MINIMAL)

colorSwitch = 0

#check to see if data is timestamp
def ifTimestamp (dataInput):
	try:
		datetime.datetime.strptime(dataInput, "%m/%d/%Y %H:%M:%S")
		return True
	except ValueError:
		return False
#end

for row in people:

	#dictionary that holds placeholder values
	ValuesDict = {'firstName': None, 'lastName': None, 'sameNumber': None, 'numberOrConfirm': None, 'roles': None, 'daysAvailable': None}

	tempList1 = []

	#ERROR HANDLING WITH DATA
	if (row[0] in (None, "")) or (ifTimestamp(row[0]) == True):
		print "ERROR: CSV file is not formatted correctly. Check included example, reformat and try again."
		sys.exit(1)
	#end

	ValuesDict.update({"firstName" : row[0]})
	ValuesDict.update({"lastName" : row[1]})
	ValuesDict.update({"sameNumber" : row[2]})
	ValuesDict.update({"numberOrConfirm" : row[3]})
	ValuesDict.update({"roles" : row[4]})
	ValuesDict.update({"daysAvailable" : row[5]})

	#check if they have a new number
	if ValuesDict.get("sameNumber") == "No":
		tempList1.append(ValuesDict.get("firstName"))
		tempList1.append(ValuesDict.get("lastName"))
		tempList1.append(ValuesDict.get("numberOrConfirm"))
		filewriter.writerow(tempList1)

	#check what roles they want and what days they can do it
	if 'Domingo (12PM)' in ValuesDict['daysAvailable']:
		if 'Pianista' in ValuesDict['roles']:
			PianoSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Bajista' in ValuesDict['roles']:
			BajoSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Guitarrista' in ValuesDict['roles']:
			GuitarraSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Apoyo' in ValuesDict['roles']:
			ApoyoSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Alabanzas' in ValuesDict['roles']:
			DirijeSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Himnos' in ValuesDict['roles']:
			HimnoSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Baterista' in ValuesDict['roles']:
			BateriaSun12.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
	#end
	if 'Domingo (6PM)' in ValuesDict['daysAvailable']:
		if 'Pianista' in ValuesDict['roles']:
			PianoSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Bajista' in ValuesDict['roles']:
			BajoSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Guitarrista' in ValuesDict['roles']:
			GuitarraSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Apoyo' in ValuesDict['roles']:
			ApoyoSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Alabanzas' in ValuesDict['roles']:
			DirijeSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Himnos' in ValuesDict['roles']:
			HimnoSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Baterista' in ValuesDict['roles']:
			BateriaSun6.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
	#end
	if 'Miercoles' in ValuesDict['daysAvailable']:
		if 'Pianista' in ValuesDict['roles']:
			PianoWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Bajista' in ValuesDict['roles']:
			BajoWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Guitarrista' in ValuesDict['roles']:
			GuitarraWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Apoyo' in ValuesDict['roles']:
			ApoyoWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Alabanzas' in ValuesDict['roles']:
			DirijeWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Dirijir Himnos' in ValuesDict['roles']:
			HimnoWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
		if 'Baterista' in ValuesDict['roles']:
			BateriaWed.append(str(ValuesDict.get("firstName")) + " " + str(ValuesDict.get("lastName")))
		#end
	#end

	#put clothing colors in dictionary
	if colorSwitch == 0:
		#print("Entered IF")
		tempVestuario.append(row[6])
		shirtColors.update({"colors" : tempVestuario})
		colorSwitch = 1
		#print(shirtColors["colors"])
	else:
		#print("Entered ELSE")
		shirtColors["colors"].append(row[6])
		#print(shirtColors["colors"])
#end

##ONLY FOR MODIFYTING THE CLOTHING COLORS LIST##

tempString = ','.join(shirtColors.get('colors'))
#print(tempString)

tempList1 = tempString.split(',')

for (i, item) in enumerate(tempList1):
	tempList1[i] = tempList1[i].replace(" ", "")
#end
#print(tempList1)

#only keep common colors
nuevoVestuario = list(set(tempList1))
#print(nuevoVestuario)

##ONLY FOR MODIFYTING THE CLOTHING COLORS LIST##

##DEBUG FOR LISTS

'''
print("HimnoWed: " + str(HimnoWed))
print("DirijeWed: " + str(DirijeWed))
print("BajoWed: " + str(BajoWed))
print("GuitarraWed: " + str(GuitarraWed))
print("PianoWed: " + str(PianoWed))
print("BateriaWed: " + str(BateriaWed))
print("ApoyoWed: " + str(ApoyoWed))

print("HimnoSun12: " + str(HimnoSun12))
print("DirijeSun12: " + str(DirijeSun12))
print("BajoSun12: " + str(BajoSun12))
print("GuitarraSun12: " + str(GuitarraSun12))
print("PianoSun12: " + str(PianoSun12))
print("BateriaSun12: " + str(BateriaSun12))
print("ApoyoSun12: " + str(ApoyoSun12))

print("HimnoSun6: " + str(HimnoSun6))
print("DirijeSun6: " + str(DirijeSun6))
print("BajoSun6: " + str(BajoSun6))
print("GuitarraSun6: " + str(GuitarraSun6))
print("PianoSun6: " + str(PianoSun6))
print("BateriaSun6: " + str(BateriaSun6))
print("ApoyoSun6: " + str(ApoyoSun6))
'''

##DEBUG FOR LISTS

#close both files
f.close()
csvfile.close()

#begins a new document
document = Document()

#line 1
paragraphL1 = document.add_paragraph("Calendario de Alabanza - " + months[monthNum.month - 1])
paragraph_format = paragraphL1.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 1.0

#line 2
paragraphL2 = document.add_paragraph("(Ensayo: Martes @8PM & Viernes @6PM & 7:15PM)")
paragraph_format = paragraphL2.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 1.0

#1x3 table with day names
table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'
row = table.rows[0]

row.cells[0].text = 'Miercoles'
changeText = row.cells[0].paragraphs[0]
run = changeText.runs
font = run[0].font
font.color.rgb = RGBColor(255, 0, 0)

row.cells[1].text = 'Domingo 12PM'
changeText = row.cells[1].paragraphs[0]
run = changeText.runs
font = run[0].font
font.color.rgb = RGBColor(255, 0, 0)

row.cells[2].text = 'Domingo 6PM'
changeText = row.cells[2].paragraphs[0]
run = changeText.runs
font = run[0].font
font.color.rgb = RGBColor(255, 0, 0)

#5x3 table with main information
table = document.add_table(rows=5, cols=3)
table.style = 'Table Grid'

row = table.rows[0]

for row in table.rows:
	for cell in row.cells:
		tr = row._tr
		trPr = tr.get_or_add_trPr()

		trHeight = OxmlElement('w:trHeight')
		trHeight.set(qn('w:val'), str(2350))
		trHeight.set(qn('w:hRule'),'atLeast')
		trPr.append(trHeight)

#main information
i = 0

for row in table.rows:
	#WEDNESDAY DATES
	if wedFlag == True:

		##DEBUG ONLY
		#print("Passed 1st WED condition.")
		#print(wednesdayDate1)
		##DEBUG ONLY

		if int(wednesdayDate1) > 3:

			if i == 0:
				row.cells[0].text = 'N/A'

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 1:
				row.cells[0].text = str(wednesdayDate1) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 2:
				row.cells[0].text = str(wednesdayDate2) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 3:
				row.cells[0].text = str(wednesdayDate3) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 4:
				row.cells[0].text = str(wednesdayDate4) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
		else:

			if i == 0:
				row.cells[0].text = str(wednesdayDate1) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 1:
				row.cells[0].text = str(wednesdayDate2) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 2:
				row.cells[0].text = str(wednesdayDate3) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 3:
				row.cells[0].text = str(wednesdayDate4) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			#end
			if i == 4:
				row.cells[0].text = str(wednesdayDate5) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[0].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
		#end

	else:

		##DEBUG ONLY
		#print("Passed 2nd WED condition.")
		#print(wednesdayDate5)
		##DEBUG ONLY

		row.cells[0].text = str(wednesdayDate1) + "-" + months[monthNum.month - 1]

		#styling text
		changeText = row.cells[0].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.color.rgb = RGBColor(0, 99, 255)
		font.size = Pt(9)

		#end
		if i == 1:
			row.cells[0].text = str(wednesdayDate2) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[0].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 2:
			row.cells[0].text = str(wednesdayDate3) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[0].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 3:
			row.cells[0].text = str(wednesdayDate4) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[0].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 4:
			row.cells[0].text = str(wednesdayDate5) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[0].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
	#end

	#SUNDAY 12PM DATES
	if sunFlag == True:

		##DEBUG ONLY
		#print("Passed 1st SUN12 condition.")
		#print(sundayDate1)
		##DEBUG ONLY

		row.cells[1].text = str(sundayDate1) + "-" + months[monthNum.month - 1]

		#styling text
		changeText = row.cells[1].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.color.rgb = RGBColor(0, 99, 255)
		font.size = Pt(9)

		if i == 1:
			row.cells[1].text = str(sundayDate2) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 2:
			row.cells[1].text = str(sundayDate3) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 3:
			row.cells[1].text = str(sundayDate4) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 4:
			if int(sundayDate5) != 0:
				row.cells[1].text = str(sundayDate5) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[1].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			else:
				row.cells[1].text = 'N/A'

				#styling text
				changeText = row.cells[1].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
		#end
	else:

		##DEBUG ONLY
		#print("Passed 2nd SUN12 condition.")
		#print(sundayDate5)
		##DEBUG ONLY

		row.cells[1].text = str(sundayDate1) + "-" + months[monthNum.month - 1]

		#styling text
		changeText = row.cells[1].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.color.rgb = RGBColor(0, 99, 255)
		font.size = Pt(9)

		if i == 1:
			row.cells[1].text = str(sundayDate2) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 2:
			row.cells[1].text = str(sundayDate3) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 3:
			row.cells[1].text = str(sundayDate4) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 4:
			row.cells[1].text = str(sundayDate5) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[1].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
	#end

	#SUNDAY 6PM DATES
	if sunFlag == True:

		##DEBUG ONLY
		#print("Passed 1st SUN6 condition.")
		#print(sundayDate1)
		##DEBUG ONLY

		row.cells[2].text = str(sundayDate1) + "-" + months[monthNum.month - 1]

		#styling text
		changeText = row.cells[2].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.color.rgb = RGBColor(0, 99, 255)
		font.size = Pt(9)

		if i == 1:
			row.cells[2].text = str(sundayDate2) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 2:
			row.cells[2].text = str(sundayDate3) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 3:
			row.cells[2].text = str(sundayDate4) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 4:
			if int(sundayDate5) != 0:
				row.cells[2].text = str(sundayDate5) + "-" + months[monthNum.month - 1]

				#styling text
				changeText = row.cells[2].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
			else:
				row.cells[2].text = 'N/A'

				#styling text
				changeText = row.cells[2].paragraphs[0]
				run = changeText.runs
				font = run[0].font
				font.color.rgb = RGBColor(0, 99, 255)
				font.size = Pt(9)
		#end
	else:

		##DEBUG ONLY
		#print("Passed 2nd SUN6 condition.")
		#print(sundayDate5)
		##DEBUG ONLY

		row.cells[2].text = str(sundayDate1) + "-" + months[monthNum.month - 1]

		#styling text
		changeText = row.cells[2].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.color.rgb = RGBColor(0, 99, 255)
		font.size = Pt(9)

		#end
		if i == 1:
			row.cells[2].text = str(sundayDate2) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 2:
			row.cells[2].text = str(sundayDate3) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 3:
			row.cells[2].text = str(sundayDate4) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
		if i == 4:
			row.cells[2].text = str(sundayDate5) + "-" + months[monthNum.month - 1]

			#styling text
			changeText = row.cells[2].paragraphs[0]
			run = changeText.runs
			font = run[0].font
			font.color.rgb = RGBColor(0, 99, 255)
			font.size = Pt(9)
		#end
	#end
	i += 1
#end

#now adding in the information from the input csv file

#BACKUP DATA TO BE SAFE
BackupDataDictionary = {
'1': HimnoWed,
'2': DirijeWed,
'3': BajoWed,
'4': GuitarraWed,
'5': PianoWed,
'6': BateriaWed,
'7': ApoyoWed,

'8': HimnoSun12,
'9': DirijeSun12,
'10': BajoSun12,
'11': GuitarraSun12,
'12': PianoSun12,
'13': BateriaSun12,
'14': ApoyoSun12,

'15': HimnoSun6,
'16': DirijeSun6,
'17': BajoSun6,
'18': GuitarraSun6,
'19': PianoSun6,
'20': BateriaSun6,
'21': ApoyoSun6
}

for x in range(1, 22):
	if not BackupDataDictionary.get(str(x)):
		BackupDataDictionary.get(str(x)).append(None)

##TEMP LISTS

tempHimnoWed = list(HimnoWed)
tempDirijeWed = list(DirijeWed)
tempBajoWed = list(BajoWed)
tempGuitarraWed = list(GuitarraWed)
tempPianoWed = list(PianoWed)
tempBateriaWed = list(BateriaWed)
tempApoyoWed = list(ApoyoWed)

tempHimnoSun12 = list(HimnoSun12)
tempDirijeSun12 = list(DirijeSun12)
tempBajoSun12 = list(BajoSun12)
tempGuitarraSun12 = list(GuitarraSun12)
tempPianoSun12 = list(PianoSun12)
tempBateriaSun12 = list(BateriaSun12)
tempApoyoSun12 = list(ApoyoSun12)

tempHimnoSun6 = list(HimnoSun6)
tempDirijeSun6 = list(DirijeSun6)
tempBajoSun6 = list(BajoSun6)
tempGuitarraSun6 = list(GuitarraSun6)
tempPianoSun6 = list(PianoSun6)
tempBateriaSun6 = list(BateriaSun6)
tempApoyoSun6 = list(ApoyoSun6)

##TEMP LISTS

#MAIN DATA INSIDE AN ENCAPSULATED DICTIONARY
MainDataDictionary = {
'1': tempHimnoWed,
'2': tempDirijeWed,
'3': tempBajoWed,
'4': tempGuitarraWed,
'5': tempPianoWed,
'6': tempBateriaWed,
'7': tempApoyoWed,

'8': tempHimnoSun12,
'9': tempDirijeSun12,
'10': tempBajoSun12,
'11': tempGuitarraSun12,
'12': tempPianoSun12,
'13': tempBateriaSun12,
'14': tempApoyoSun12,

'15': tempHimnoSun6,
'16': tempDirijeSun6,
'17': tempBajoSun6,
'18': tempGuitarraSun6,
'19': tempPianoSun6,
'20': tempBateriaSun6,
'21': tempApoyoSun6
}

def dataChooser (serviceInput):
	if len(MainDataDictionary.get(str(serviceInput))) == 0:

		newList = list(BackupDataDictionary.get(str(serviceInput)))
		MainDataDictionary.update({str(serviceInput): newList})
		randomIndex = random.randint(0,(len(newList) - 1))
		nameSelected = MainDataDictionary.get(str(serviceInput))[randomIndex]
		MainDataDictionary.get(str(serviceInput)).remove(nameSelected)

		return nameSelected
	else:

		randomIndex = random.randint(0,(len(MainDataDictionary.get(str(serviceInput))) - 1))
		nameSelected = MainDataDictionary.get(str(serviceInput))[randomIndex]
		MainDataDictionary.get(str(serviceInput)).remove(nameSelected)

		return nameSelected
#end

tempnuevoVestuario = list(nuevoVestuario)

def colorChooser (tempnuevoVestuario):
	if not tempnuevoVestuario:
		tempnuevoVestuario = list(nuevoVestuario)
		randomIndex = random.randint(0,(len(tempnuevoVestuario) - 1))
		colorSelected = tempnuevoVestuario[randomIndex]
		tempnuevoVestuario.remove(colorSelected)

		return colorSelected
	else:
		randomIndex = random.randint(0,(len(tempnuevoVestuario) - 1))
		colorSelected = tempnuevoVestuario[randomIndex]
		tempnuevoVestuario.remove(colorSelected)

		return colorSelected
#end

for row in table.rows:
	if row.cells[0].text == 'N/A':
		#styling text
		changeText = row.cells[0].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)
	else:
		serviceInput = 1
		row.cells[0].add_paragraph("Himno: " + str(dataChooser(serviceInput)))
		serviceInput = 2
		row.cells[0].add_paragraph("Dirije: " + str(dataChooser(serviceInput)))
		serviceInput = 3
		row.cells[0].add_paragraph("Bajo: " + str(dataChooser(serviceInput)))
		serviceInput = 4
		row.cells[0].add_paragraph("Guitarra: " + str(dataChooser(serviceInput)))
		serviceInput = 5
		row.cells[0].add_paragraph("Piano: " + str(dataChooser(serviceInput)))
		serviceInput = 6
		row.cells[0].add_paragraph("Bateria: " + str(dataChooser(serviceInput)))
		serviceInput = 7
		row.cells[0].add_paragraph("Apoyo: " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)))
		row.cells[0].add_paragraph("Vestuario: " + str(colorChooser(tempnuevoVestuario)) + " y Negro")

		#styling text
		changeText = row.cells[0].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

		for x in range(1, 8):
			changeText = row.cells[0].paragraphs[x]
			run = changeText.runs
			font = run[0].font
			font.size = Pt(9)
		#end

		#styling text
		changeText = row.cells[0].paragraphs[8]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

	#end
	if row.cells[1].text == 'N/A':
		#styling text
		changeText = row.cells[1].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)
	else:
		serviceInput = 8
		row.cells[1].add_paragraph("Himno: " + str(dataChooser(serviceInput)))
		serviceInput = 9
		row.cells[1].add_paragraph("Dirije: " + str(dataChooser(serviceInput)))
		serviceInput = 10
		row.cells[1].add_paragraph("Bajo: " + str(dataChooser(serviceInput)))
		serviceInput = 11
		row.cells[1].add_paragraph("Guitarra: " + str(dataChooser(serviceInput)))
		serviceInput = 12
		row.cells[1].add_paragraph("Piano: " + str(dataChooser(serviceInput)))
		serviceInput = 13
		row.cells[1].add_paragraph("Bateria: " + str(dataChooser(serviceInput)))
		serviceInput = 14
		row.cells[1].add_paragraph("Apoyo: " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)))
		row.cells[1].add_paragraph("Vestuario: " + str(colorChooser(tempnuevoVestuario)) + " y Negro")

		#styling text
		changeText = row.cells[1].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

		for x in range(1, 8):
			changeText = row.cells[1].paragraphs[x]
			run = changeText.runs
			font = run[0].font
			font.size = Pt(9)
		#end

		#styling text
		changeText = row.cells[1].paragraphs[8]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

	#end
	if row.cells[2].text == 'N/A':
		#styling text
		changeText = row.cells[2].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)
	else:
		serviceInput = 15
		row.cells[2].add_paragraph("Himno: " + str(dataChooser(serviceInput)))
		serviceInput = 16
		row.cells[2].add_paragraph("Dirije: " + str(dataChooser(serviceInput)))
		serviceInput = 17
		row.cells[2].add_paragraph("Bajo: " + str(dataChooser(serviceInput)))
		serviceInput = 18
		row.cells[2].add_paragraph("Guitarra: " + str(dataChooser(serviceInput)))
		serviceInput = 19
		row.cells[2].add_paragraph("Piano: " + str(dataChooser(serviceInput)))
		serviceInput = 20
		row.cells[2].add_paragraph("Bateria: " + str(dataChooser(serviceInput)))
		serviceInput = 21
		row.cells[2].add_paragraph("Apoyo: " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)) + ", " + str(dataChooser(serviceInput)))
		row.cells[2].add_paragraph("Vestuario: " + str(colorChooser(tempnuevoVestuario)) + " y Negro")

		#styling text
		changeText = row.cells[2].paragraphs[0]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

		for x in range(1, 8):
			changeText = row.cells[2].paragraphs[x]
			run = changeText.runs
			font = run[0].font
			font.size = Pt(9)
		#end

		#styling text
		changeText = row.cells[2].paragraphs[8]
		run = changeText.runs
		font = run[0].font
		font.bold = True
		font.size = Pt(9)

#last line
paragraphLL = document.add_paragraph()
paragraph_format = paragraphLL.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 1.0

change = paragraphLL.add_run("ESTAREMOS LIDERANDO EL SERVICIO DE ORACION EL MARTES " + str(tuesdayDate) + " DE " + months[monthNum.month - 1].upper() + ".")
change.font.italic = True
change.font.bold = True
change.font.color.rgb = RGBColor(255, 0, 0)
change.font.highlight_color = WD_COLOR_INDEX.YELLOW

#saves document in same directory as program file
document.save('Mes de ' + str(months[monthNum.month - 1]) + '.docx')